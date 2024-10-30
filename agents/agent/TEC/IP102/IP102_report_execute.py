import json
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import sys

def read_json(json_file_path):
    try:
        with open(json_file_path, 'r') as file:
            data = json.load(file)
        return data
    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
    except FileNotFoundError:
        print(f"File not found: {json_file_path}")
    except Exception as e:
        print(f"An error occurred: {e}")

def replace_placeholders_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, value)

def replace_placeholders_in_docx(doc, replacements):
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)
    
    return doc

def add_images_to_docx_at_words(doc, base_image_path, image_data):
    for item in image_data:
        word = item["word"]
        image_path = os.path.join(base_image_path, item["image_path"])
        caption = item.get("caption", "")

        if not os.path.exists(image_path):
            print(f"Image not found: {image_path}")
            continue

        for paragraph in doc.paragraphs:
            if word in paragraph.text:
                try:
                    img_paragraph = paragraph.insert_paragraph_before()
                    run = img_paragraph.add_run()
                    run.add_picture(image_path, width=Inches(5.9))
                    img_paragraph.alignment = 1

                    if caption:
                        caption_paragraph = paragraph.insert_paragraph_before()
                        caption_run = caption_paragraph.add_run(caption)
                        caption_run.font.name = 'Calibri'
                        caption_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
                        caption_run.font.size = Pt(10)
                        caption_run.font.bold = True
                        caption_run.font.italic = True
                        caption_paragraph.alignment = 1

                    paragraph.text = paragraph.text.replace(word, "")
                except Exception as e:
                    print(f"Failed to add image '{image_path}' for word '{word}': {e}")

    return doc


def main():
    if len(sys.argv) != 4:
        print("Usage: python3 report.py <scan_id> <result_path> <folder_path>")
        sys.exit(1)

    # Extract arguments from the command-line
    scan_id = sys.argv[1]
    result_path = sys.argv[2]
    folder_path = sys.argv[3]

    # Construct the path to the scan-specific JSON file
    json_file_path = os.path.join(result_path, "dut_conf", "dut_details.json")
    
    # Use folder_path for the DOCX template
    docx_path = os.path.join(folder_path, "IP102_Template.docx") 
    output_path = os.path.join(result_path, f"{scan_id}.docx") 
    
    data = read_json(json_file_path)
    if data is not None:
        replacements = data.get("placeholders", {})
        base_image_path = ""
        image_data = data.get("images", [])
        
        try:
            doc = Document(docx_path)
        except Exception as e:
            print(f"Failed to load the document: {e}")
            doc = None

        if doc is not None:
            doc = replace_placeholders_in_docx(doc, replacements)
            doc = add_images_to_docx_at_words(doc, base_image_path, image_data)

            try:
                doc.save(output_path)
                print(f"Document saved successfully to {output_path}")
            except Exception as e:
                print(f"Failed to save the document: {e}")
        else:
            print("Failed to update the DOCX file with placeholders.")
    else:
        print("Failed to read JSON data.")

if __name__ == "__main__":
    main()